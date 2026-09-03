"""PDF intelligence & forensics (spec #19, #20).

Extraction: text (pypdf, PyMuPDF fallback), metadata, page count, embedded
images, links/annotations, headings, tables (PyMuPDF).

Forensics heuristics (never claim fraud from metadata alone — spec #20):
  * creation vs modification date contradictions
  * producer/creator tool anomalies (converter chains, image-only "PDFs")
  * XMP vs Info-dict date mismatches
  * empty-text pages (scanned or image-only content)
  * suspicious link patterns
"""
from __future__ import annotations

import io
import re
from datetime import datetime, timezone
from typing import Any, Optional

from app.core.logging import get_logger

logger = get_logger("prvision.media.pdf")

try:
    import fitz  # PyMuPDF
    HAS_FITZ = True
except Exception:  # pragma: no cover
    HAS_FITZ = False

from pypdf import PdfReader


def _pdf_date(value: Optional[str]) -> Optional[datetime]:
    if not value:
        return None
    v = str(value).strip()
    m = re.match(r"[Dd]:?(\d{4})(\d{2})?(\d{2})?(\d{2})?(\d{2})?(\d{2})?", v)
    if not m:
        try:
            return datetime.fromisoformat(v.replace("Z", "+00:00"))
        except ValueError:
            return None
    try:
        return datetime(
            int(m.group(1)), int(m.group(2) or 1), int(m.group(3) or 1),
            int(m.group(4) or 0), int(m.group(5) or 0), int(m.group(6) or 0),
            tzinfo=timezone.utc,
        )
    except ValueError:
        return None


def extract_pdf(data: bytes) -> dict[str, Any]:
    """Extract text + metadata + structure. Returns parser-agnostic dict."""
    out: dict[str, Any] = {
        "text": "", "title": None, "author": None, "subject": None,
        "creator": None, "producer": None, "creation_date": None, "mod_date": None,
        "pages": 0, "images": 0, "links": [], "headings": [],
        "xmp_dates": {}, "empty_text_pages": 0, "forensics": [],
    }

    text_parts: list[str] = []
    try:
        reader = PdfReader(io.BytesIO(data))
        out["pages"] = len(reader.pages)
        meta = reader.metadata or {}
        out["title"] = (meta.get("/Title") or None)
        out["author"] = (meta.get("/Author") or None)
        out["subject"] = (meta.get("/Subject") or None)
        out["creator"] = (meta.get("/Creator") or None)
        out["producer"] = (meta.get("/Producer") or None)
        out["creation_date"] = _pdf_date(meta.get("/CreationDate"))
        out["mod_date"] = _pdf_date(meta.get("/ModDate"))
        for page in reader.pages:
            try:
                text_parts.append(page.extract_text() or "")
            except Exception:
                text_parts.append("")
        for page in reader.pages:
            for annot in (page.get("/Annots") or []):
                try:
                    obj = annot.get_object()
                    uri = obj.get("/A", {}).get("/URI")
                    if uri:
                        out["links"].append(str(uri))
                except Exception:
                    continue
    except Exception as exc:
        logger.warning("pypdf extraction limited: %s", exc.__class__.__name__)

    # PyMuPDF enrichment: images, links, headings, per-page empties
    if HAS_FITZ:
        try:
            doc = fitz.open(stream=data, filetype="pdf")
            out["pages"] = max(out["pages"], doc.page_count)
            images = 0
            empty_pages = 0
            links: list[str] = list(out["links"])
            headings: list[str] = []
            for page in doc:
                images += len(page.get_images(full=True))
                ptext = page.get_text("text").strip()
                if len(ptext) < 20:
                    empty_pages += 1
                for blk in page.get_text("dict").get("blocks", []):
                    for line in blk.get("lines", []):
                        for span in line.get("spans", []):
                            if span.get("size", 0) >= 15 and span.get("text", "").strip():
                                headings.append(span["text"].strip()[:160])
                links += [l.get("uri") for l in page.get_links() if l.get("uri")]
            out["images"] = images
            out["empty_text_pages"] = empty_pages
            out["links"] = list(dict.fromkeys(links))[:60]
            out["headings"] = list(dict.fromkeys(headings))[:40]
            # XMP dates
            try:
                xmp = doc.xmp_metadata
                if xmp:
                    for key in ("xmp:CreateDate", "xmp:ModifyDate", "pdf:Producer"):
                        val = getattr(xmp, key.split(":")[1] if False else key.replace(":", "_"), None)
                        if val:
                            out["xmp_dates"][key] = str(val)
            except Exception:
                pass
            doc.close()
        except Exception as exc:
            logger.warning("PyMuPDF enrichment failed: %s", exc)

    text = "\n\n".join(t for t in text_parts if t)
    out["text"] = text[:400_000]
    out["forensics"] = pdf_forensics(out)
    return out


def pdf_forensics(parsed: dict[str, Any]) -> list[dict[str, str]]:
    """Heuristic document forensics (spec #20). Every finding is framed as a
    signal to review — metadata anomalies are NOT proof of fraud."""
    findings: list[dict[str, str]] = []
    created, modified = parsed.get("creation_date"), parsed.get("mod_date")

    if created and modified:
        delta_days = (modified - created).days
        if delta_days > 30:
            findings.append({
                "signal": "late_modification",
                "detail": f"Modified {delta_days} days after creation — document was updated post-issue; verify which version is authoritative.",
            })
        if modified < created:
            findings.append({
                "signal": "date_contradiction",
                "detail": "Modification date precedes creation date — timestamp inconsistency worth reviewing.",
            })

    xmp = parsed.get("xmp_dates") or {}
    if created and xmp:
        for key, raw in xmp.items():
            xmp_dt = _pdf_date(str(raw))
            if xmp_dt and created and abs((xmp_dt - created).days) > 7:
                findings.append({
                    "signal": "xmp_info_mismatch",
                    "detail": f"{key} disagrees with the Info-dict creation date — re-saved through a different tool?",
                })

    producer = (parsed.get("producer") or "").lower()
    creator = (parsed.get("creator") or "").lower()
    if producer and not re.search(r"(acrobat|word|indesign|latex|libreoffice|openoffice|pages|quartz|preview|pandoc)", producer + " " + creator):
        findings.append({
            "signal": "unusual_producer",
            "detail": f"Unusual producer tool “{parsed.get('producer')}” — common with converters; not proof of tampering.",
        })

    if parsed.get("pages") and parsed.get("empty_text_pages", 0) >= max(1, int(parsed["pages"] * 0.6)):
        findings.append({
            "signal": "image_only_pages",
            "detail": "Most pages contain no extractable text — likely scans or image-only export; OCR verification recommended.",
        })

    links = parsed.get("links") or []
    shorteners = sum(1 for l in links if re.search(r"(bit\.ly|tinyurl|t\.co|goo\.gl|is\.gd|shorturl)", (l or ""), re.IGNORECASE))
    if links and shorteners / len(links) > 0.3:
        findings.append({
            "signal": "suspicious_links",
            "detail": f"{shorteners}/{len(links)} embedded links are URL shorteners — common in spam/phishing documents.",
        })

    text = parsed.get("text") or ""
    if text:
        # repeated identical sections (template spam / duplicated appendices)
        chunks = [c.strip() for c in re.split(r"\n{2,}", text) if len(c.strip()) > 120]
        seen: dict[str, int] = {}
        repeats = 0
        for c in chunks:
            key = c[:80]
            seen[key] = seen.get(key, 0) + 1
            if seen[key] == 2:
                repeats += 1
        if repeats >= 3:
            findings.append({
                "signal": "repeated_sections",
                "detail": f"{repeats} identical long sections repeated — duplicated content worth reviewing.",
            })

    return findings


# --------------------------------------------------------------------------- DOCX
def extract_docx(data: bytes) -> dict[str, Any]:
    """DOCX text + core properties (spec #21)."""
    import io as _io
    from docx import Document

    out: dict[str, Any] = {"text": "", "title": None, "author": None,
                           "created": None, "modified": None, "paragraphs": 0,
                           "tables": 0, "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
    try:
        doc = Document(_io.BytesIO(data))
        paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        out["paragraphs"] = len(paras)
        out["tables"] = len(doc.tables)
        core = doc.core_properties
        out["title"] = core.title or None
        out["author"] = core.author or None
        out["created"] = core.created.isoformat() if core.created else None
        out["modified"] = core.modified.isoformat() if core.modified else None
        # include table text (often where claims live)
        table_lines: list[str] = []
        for t in doc.tables[:20]:
            for row in t.rows[:60]:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    table_lines.append(" | ".join(cells))
        out["text"] = "\n".join(paras + ["", "TABLES:"] + table_lines)[:300_000]
    except Exception as exc:
        logger.warning("docx parse failed: %s", exc.__class__.__name__)
        out["text"] = ""
    return out

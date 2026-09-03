"""DOCX intelligence — text, tables, metadata, links (spec #8).

Dedicated module (fixes the runtime bug where verification/ingestion.py
imported app.media.docx_analysis but only a pdf_analysis.extract_docx
re-export existed, so every DOCX upload crashed with ModuleNotFoundError).

Extraction is deterministic and read-only: python-docx for body/tables/core
properties, relationship scan for hyperlinks. Nothing here executes
document content — documents are untrusted data.
"""
from __future__ import annotations

import hashlib
import io
import zipfile
from typing import Any

from app.core.logging import get_logger

logger = get_logger("prvision.media.docx")

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _is_docx(data: bytes) -> bool:
    """DOCX is a ZIP container starting with the PK\x03\x04 local-file header."""
    return len(data) > 4 and data[:2] == b"PK"


def extract_docx(data: bytes, filename: str = "document.docx") -> dict[str, Any]:
    """Extract text, tables, metadata and links from DOCX bytes.

    Returns a dict; never raises for malformed documents — returns a
    structured error payload instead so the pipeline can degrade honestly.
    """
    result: dict[str, Any] = {
        "mime": DOCX_MIME,
        "size": len(data),
        "sha256": hashlib.sha256(data).hexdigest(),
        "title": None,
        "author": None,
        "created": None,
        "modified": None,
        "text": "",
        "paragraphs": 0,
        "tables": [],
        "links": [],
        "images": 0,
        "forensics": {},
    }
    if not _is_docx(data):
        result["forensics"] = {"error": "not_a_zip_container",
                               "note": "File does not have the DOCX (ZIP) structure — "
                                       "it may be corrupted, renamed, or a legacy .doc binary."}
        return result

    try:
        import docx  # python-docx
    except ImportError:  # pragma: no cover - dependency present in requirements
        result["forensics"] = {"error": "docx_library_unavailable"}
        return result

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # malformed / password-protected / not real docx
        logger.warning("DOCX parse failed for %s: %s", filename, exc)
        result["forensics"] = {"error": "docx_parse_failed", "detail": exc.__class__.__name__,
                               "note": "Document could not be opened — it may be corrupted "
                                       "or not a genuine DOCX file."}
        return result

    # ---- core properties (metadata forensics)
    try:
        props = document.core_properties
        result["title"] = (props.title or "").strip() or None
        result["author"] = (props.author or "").strip() or None
        result["created"] = props.created.isoformat() if props.created else None
        result["modified"] = props.modified.isoformat() if props.modified else None
        last_mod = (props.last_modified_by or "").strip()
        if last_mod:
            result["forensics"]["last_modified_by"] = last_mod[:200]
        rev = props.revision
        if isinstance(rev, int) and rev > 0:
            result["forensics"]["revision_count"] = rev
        if result["created"] and result["modified"] and result["modified"] < result["created"]:
            result["forensics"]["anomalies"] = result["forensics"].get("anomalies", [])
            result["forensics"]["anomalies"].append("modified_precedes_created")
    except Exception:  # properties are optional
        pass

    # ---- body text
    paragraphs = [p.text.strip() for p in document.paragraphs]
    text = "\n\n".join(p for p in paragraphs if p)
    result["paragraphs"] = sum(1 for p in paragraphs if p)
    result["text"] = text[:200_000]

    # ---- tables (numeric verification downstream uses these)
    tables_out: list[dict[str, Any]] = []
    for t_idx, table in enumerate(document.tables[:25]):
        rows = []
        for row in table.rows[:60]:
            rows.append([cell.text.strip()[:200] for cell in row.cells[:25]])
        if not rows:
            continue
        header = rows[0]
        numeric_columns: list[str] = []
        for col_idx, col in enumerate(header):
            vals = [r[col_idx] for r in rows[1:11] if col_idx < len(r) and r[col_idx].strip()]
            ok = 0
            for v in vals:
                try:
                    float(v.replace(",", "").replace("%", "").replace(" ", ""))
                    ok += 1
                except ValueError:
                    pass
            if vals and ok / len(vals) > 0.7:
                numeric_columns.append(col)
        tables_out.append({"index": t_idx, "rows": len(rows), "columns": len(header),
                           "header": header, "sample_rows": rows[1:6],
                           "numeric_columns": numeric_columns})
    result["tables"] = tables_out

    # ---- hyperlinks (via document part relationships)
    try:
        rels = document.part.rels
        links = []
        for rel in rels.values():
            if "hyperlink" in rel.reltype and rel.target_ref:
                links.append(str(rel.target_ref)[:500])
        result["links"] = links[:100]
    except Exception:
        pass

    # ---- embedded media count (zip scan; cheap + robust)
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            result["images"] = sum(1 for n in zf.namelist() if n.startswith("word/media/"))
    except zipfile.BadZipFile:
        pass

    return result

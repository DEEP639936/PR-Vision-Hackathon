"""Unit tests for the security primitives (spec #17)."""
from __future__ import annotations

import pytest


# ------------------------------------------------------------------ passwords
def test_password_hash_roundtrip():
    from app.core.security import hash_password, verify_password

    stored = hash_password("Sup3rSecret!")
    assert stored.startswith("pbkdf2_sha256$")
    assert "Sup3rSecret!" not in stored
    assert verify_password("Sup3rSecret!", stored) is True
    assert verify_password("wrong", stored) is False
    assert verify_password("", stored) is False
    # unique salts per call
    assert hash_password("Sup3rSecret!") != stored


def test_password_policy():
    from app.core.security import password_issues

    assert password_issues("Sup3rSecret!") == []
    assert len(password_issues("short1")) >= 1
    assert password_issues("nodigits!") != []
    assert password_issues("123456789") != []


# ------------------------------------------------------------------ SSRF guard
@pytest.mark.parametrize("host,allowed", [
    ("1.1.1.1", True),
    ("8.8.8.8", True),
    ("localhost", False),
    ("127.0.0.1", False),
    ("0.0.0.0", False),
    ("10.0.0.1", False),
    ("169.254.169.254", False),      # cloud metadata endpoint
    ("100.64.0.1", False),           # CGNAT
    ("192.168.1.1", False),
    ("172.16.0.1", False),
    ("172.32.0.1", True),            # outside the private 172.16/12 range
    ("224.0.0.1", False),            # multicast
    ("::1", False),
    ("::ffff:10.0.0.1", False),      # IPv6-mapped IPv4
    ("fd00::1", False),              # IPv6 ULA
    ("fe80::1", False),              # IPv6 link-local
])
def test_ssrf_host_validation(host, allowed):
    from app.verification.ingestion import _validate_resolved_ip

    assert _validate_resolved_ip(host) is allowed, host


def test_validate_url_blocks_credentials():
    from app.verification.ingestion import validate_url, IngestionError

    with pytest.raises(IngestionError):
        validate_url("https://user:pass@example.com/page")


# ------------------------------------------------------------------ rate limiter
def _make_client(monkeypatch, limit: int, enabled: bool = True):
    """Standalone app with the rate limiter patched to a tiny auth limit."""
    from fastapi import FastAPI
    from fastapi.testclient import TestClient

    from app.core.config import settings
    from app.core.ratelimit import RateLimitMiddleware

    monkeypatch.setattr(settings, "RATE_LIMIT_AUTH_PER_MINUTE", limit)
    monkeypatch.setattr(settings, "RATE_LIMIT_ENABLED", enabled)

    test_app = FastAPI()
    test_app.add_middleware(RateLimitMiddleware)

    @test_app.get("/api/auth/ping")
    def ping():
        return {"ok": True}

    @test_app.get("/api/health")
    def health():
        return {"ok": True}

    return TestClient(test_app)


def test_rate_limiter_blocks_after_limit(monkeypatch):
    client = _make_client(monkeypatch, 3)
    for _ in range(3):
        assert client.get("/api/auth/ping").status_code == 200
    r = client.get("/api/auth/ping")
    assert r.status_code == 429
    assert int(r.headers["retry-after"]) >= 1
    # health is exempt
    assert client.get("/api/health").status_code == 200


def test_rate_limiter_can_be_disabled(monkeypatch):
    client = _make_client(monkeypatch, 1, enabled=False)
    for _ in range(3):
        assert client.get("/api/auth/ping").status_code == 200


# ------------------------------------------------------------------ alert engine
def test_alert_severity_mapping():
    from app.services.alert_engine import _severity_from_risk

    assert _severity_from_risk(0.95) == "CRITICAL"
    assert _severity_from_risk(0.8) == "HIGH"
    assert _severity_from_risk(0.55) == "MEDIUM"
    assert _severity_from_risk(0.1) == "LOW"


def test_alert_dedupe_suppresses_repeats(db_session):
    from app.db.models import Alert
    from app.db.repositories import AlertRepository

    AlertRepository.add(db_session, kind="acceleration_spike", severity="HIGH",
                        title="t", message="m", metrics={}, dedupe_key="unit:dedupe")
    db_session.flush()
    assert AlertRepository.recent_duplicate(db_session, "unit:dedupe") is not None
    assert AlertRepository.recent_duplicate(db_session, "other:key") is None
    count_before = db_session.query(Alert).count()
    # engine must NOT raise a duplicate within the window
    from app.services.alert_engine import _raise

    raised = _raise(db_session, kind="acceleration_spike", severity="HIGH",
                    title="t2", message="m2", metrics={}, dedupe_key="unit:dedupe")
    assert raised is False
    assert db_session.query(Alert).count() == count_before


# ------------------------------------------------------------------ DOCX module
def test_docx_module_rejects_garbage_and_accepts_real():
    from app.media.docx_analysis import extract_docx

    bad = extract_docx(b"this is not a docx at all")
    assert bad["forensics"]["error"] in ("not_a_zip_container", "docx_parse_failed")

    # build a minimal real DOCX with python-docx
    import io

    import docx as docx_lib

    document = docx_lib.Document()
    document.add_paragraph("The water plant releases 40% more contaminants.")
    document.add_paragraph("Officials deny the claim.")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "year"
    table.rows[0].cells[1].text = "value"
    table.rows[1].cells[0].text = "2024"
    table.rows[1].cells[1].text = "42"
    buf = io.BytesIO()
    document.save(buf)
    parsed = extract_docx(buf.getvalue())
    assert parsed["paragraphs"] == 2
    assert "40% more contaminants" in parsed["text"]
    assert parsed["tables"], "table not extracted"
    # both columns parse as numeric (year + value are all-numeric samples)
    assert set(parsed["tables"][0]["numeric_columns"]) == {"year", "value"}
